import streamlit as st
import hmac
import io
import os
import openai
from openai import OpenAI
from docx import Document, table
#from langchain.prompts import PromptTemplate
#from langchain_community.embeddings import OpenAIEmbeddings
import docx
from docx.shared import Inches, Pt, Cm
import PyPDF2
#import pdfplumber
import openpyxl

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
#from dotenv import load_dotenv


#import docx2txt

def extract_paragraphs_by_header(source_document, target_document,header_style, target_words):
    extracted_paragraphs = []
    below_header = False
    for paragraph in source_document.paragraphs:
        if paragraph.style.name == header_style:
            below_header = True
            extracted_paragraphs.append(paragraph.text)
            copy_table_data(source_document, target_document)
        elif below_header:
            extracted_paragraphs.append(paragraph.text)
            # Check if the header text contains the target words
            if all(word.lower() in paragraph.text.lower() for word in target_words):
                break  # Exit the loop if found
    return extracted_paragraphs


def copy_table_data(source_document, target_document):
  
  # Iterate through tables in the source document
  for source_table in source_document.tables:
    # Create a new empty table in the target document with the same dimensions
    target_table = target_document.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    num_rows = len(target_table.rows)

    # Iterate through rows and cells in the source table
    for source_row_index, source_row in enumerate(source_table.rows):
      for source_cell_index, source_cell in enumerate(source_row.cells):
        # Copy cell text from the source table to the target table
        target_table.cell(source_row_index, source_cell_index).text = source_cell.text

def merge_components(merged_document, components):
    for component in components:
        if isinstance(component, docx.text.paragraph.Paragraph):
            merged_document.add_paragraph(component.text)
        elif isinstance(component, docx.table.Table):
            merged_document.add_table(component)
        elif isinstance(component, docx.shape.InlineShape):
            merged_document.add_picture(component.part.blob, width=Inches(2))  # Adjust image size as needed
    return merged_document

def save_document(merged_document, file_name):
    merged_document.save(file_name)

def process_documents(documents, output_file):
    extracted_paragraphs_list = []
    for document in documents:
        if document:
            extracted_paragraphs = extract_paragraphs_by_header(document, "Heading 1")
            extracted_paragraphs_list.append(extracted_paragraphs)

    # Create a new Word document and save the extracted paragraphs
    new_document = Document()
    for paragraphs in extracted_paragraphs_list:
        for paragraph in paragraphs:
            new_document.add_paragraph(paragraph)
            for tables in extracted_tables:
                for table in tables:
                    new_document.add_table(table)
    new_document.save(output_file)

def extract_text(pdf_file):
    with open(pdf_file, 'rb') as pdf_reader:
        pdf_reader = PyPDF2.PdfReader(pdf_reader)
        text = ''
        for pagenum in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[pagenum]
            text += page.extract_text()
        return text

#used to read text in a pdf file. Use this
def read_pdf_file(file_path):
  """Reads a PDF file and returns its text content.

  Args:
    file_path: The path to the PDF file.

  Returns:
    The text content of the PDF file.
  """

  pdffile = io.BytesIO(pdf_bytes)
  pdf_reader = PyPDF2.PdfReader(pdffile)  # No need for 'open' here

  text = ''
  for page_num in range(len(pdf_reader.pages)):
      page = pdf_reader.pages[page_num]
      text += page.extract_text().replace('\n', '')

  return text  


def read_excel_data(filename):
    """Reads data from an Excel file and returns a list of lists.

    Args:
        filename: The path to the Excel file.

    Returns:
        A list of lists, where each inner list represents a row of data.
    """

    try:
        workbook = openpyxl.load_workbook(filename, data_only=True)
        sheet = workbook.active  # Assuming the first sheet is active

        data = []
        for row in sheet.iter_rows(values_only=True):
            data.append(list(row))

        return data

    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return None

#for login segment
def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if hmac.compare_digest(st.session_state["password"], st.secrets["password"]):
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Don't store the password.
        else:
            st.session_state["password_correct"] = False

    # Return True if the password is validated.
    if st.session_state.get("password_correct", False):
        return True

    # Show input for password.
    st.text_input(
        "Password", type="password", on_change=password_entered, key="password"
    )
    if "password_correct" in st.session_state:
        st.error("😕 Password incorrect")
    return False

 #This is the "Updated" helper function for calling LLM

def get_completion2(prompt, model="gpt-4o-mini", temperature=0, top_p=1.0, max_tokens=1024, n=1, json_output=False):
    if json_output == True:
      output_json_structure = {"type": "json_object"}
    else:
      output_json_structure = None

   # Update the prompt with extracted text
    
    response = client.chat.completions.create( #originally was openai.chat.completions
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
        n=1,
        response_format=output_json_structure,
    )
    return response.choices[0].message.content

def get_completion(prompt, model="gpt-4o-mini", temperature=0, top_p=1.0, max_tokens=256, n=1, json_output=False):
    if json_output == True:
      output_json_structure = {"type": "json_object"}
    else:
      output_json_structure = None

    response = client.chat.completions.create( #originally was openai.chat.completions
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
        n=1,
        response_format=output_json_structure,
    )
    return response.choices[0].message.content

def chatgpt_response_to_list(response_text):
  """
  Converts a ChatGPT response text into a list.

  Args:
    response_text (str): The text of the ChatGPT response.

  Returns:
    list: A list containing the elements of the response.
  """

  # Attempt to split the response by newlines, semicolons, or commas
  potential_delimiters = ["\n", ";", ","]
  for delimiter in potential_delimiters:
    response_list = response_text.split(delimiter)
    if len(response_list) > 1:
      return response_list

  # If no delimiters are found, return the entire response as a single-element list
  return [response_text]

#set bold for document heading
def set_all_headings_bold(document):
    for style in document.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith('Heading'):
            style.font.bold = True
            style.font.color.rgb = RGBColor(0,0,0)
            # You might want to set sizes here based on the heading level
            if style.name == 'Heading 1':
                style.font.size = Pt(12)
            elif style.name == 'Heading 2':
                style.font.size = Pt(14)
            elif style.name == 'Heading 3':
                style.font.size = Pt(12)
            
#set font for document.
def set_document_font(document, font_name, font_size):
    # Get the default style
    style = document.styles['Normal']
    
    # Set the font name
    style.font.name = font_name
    
    # Set the font size
    style.font.size = Pt(font_size)
    
    # Apply the style to all paragraphs in the document
    for paragraph in document.paragraphs:
        if paragraph.style.name == 'Heading 1':
            paragraph.style = style
        paragraph.style = style
        
    # Apply the style to all table cells in the document
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = style

def set_heading_styles(document, font_name, base_size): 
    for i in range(1, 4): 
        # Adjust headings 1-3 
        style = document.styles[f'Heading {i}'] 
        style.font.name = font_name 
        style.font.size = Pt(base_size + 4 - i) # Larger size for higher-level headings 
        style.font.bold = True

#set background color for cell
def set_cell_background(cell, fill):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_colour(table):
    # Set the header row background to grey
    header_row = table.rows[0]
    for cell in header_row.cells:
        set_cell_background(cell, "D3D3D3")  # Light grey

    # Optionally, make header text bold
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

def force_font(text):
    heading_1= document.add_heading(text, level=1)
    for run in heading_1.runs:
        run.font.name = 'Arial'  # Change the font to Arial
        run.font.size = Pt(12)   # Optional: Change font size

def round_to_nearest_thousand(number):
  """Rounds a number to the nearest thousand.

  Args:
    number: The number to round.

  Returns:
    The rounded number.
  """
  rounded_number = round(number / 1000.0) * 1000
  formatted_number = f"{rounded_number:,}"
  return formatted_number

def check_openai_key(api_key):
    try:
        # Set the API key
        openai.api_key = api_key

        # Make a simple request to the OpenAI API (e.g., list models)
        openai.Model.list()

        # If the request is successful, the key is valid
        print("API key is valid.")
        return True

    except openai.AuthenticationError:
        print("API key is invalid.")
        return False

    except Exception as e:
        print(f"An error occurred: {e}")
        return False

#if not check_password():
#    st.stop()  # Do not continue if check_password is not True.

#instructions for the app
st.title("Welcome to the LeverAIge Award Recommendation Assistant!🤖")
st.sidebar.title("About this App🤖")
st.sidebar.write("This app will help you create a new draft of an Award Recommendation Document.")

st.sidebar.title("Instructions:")
instructions = [
    "Upload Documents: Drag and drop or click the \"Upload Documents\" button to select the three required documents.",
    "Let GenAI do the work, and your draft Award Recommendation is completed!"
]

for i, instruction in enumerate(instructions, start=1):
    st.sidebar.markdown(f"{i}. {instruction}")

st.sidebar.title("Contact us:")
st.sidebar.write("Please reach out to us at XXX if you have any queries.")
st.sidebar.write("Made for the CPF Community by Barry Ng, Hazel Ong and Lee Puay Tiang📱")
st.sidebar.title("Important Notice:")
st.sidebar.write("This web application is developed as a proof-of-concept prototype. The information provided here is NOT intended for actual usage and should not be relied upon for making any decisions, especially those related to financial, legal, or healthcare matters.")
st.sidebar.write("Furthermore, please be aware that the LLM may generate inaccurate or incorrect information. You assume full responsibility for how you use any generated output.")
st.sidebar.write("Always consult with qualified professionals for accurate and personalized advice.")
#options = ['A', 'B', 'C']
#selected_option = st.selectbox('Select a Tender Board', options)
selected_option = "A"

#documents = load_documents()

#to obtain API key
#load_dotenv()
API_KEY = st.text_input("Please input your OpenAI API Key", type="password")

if not API_KEY.startswith('sk-'):
    st.warning("Please enter a valid API Key")
    st.stop()

client = OpenAI(api_key=API_KEY)
#AOR Document Processing
#doc = docx.Document("AOR_Sample_Desen.docx")
#Evaluation document Processing Excel
#data = read_excel_data('Detailed Eval Report_Desen.xlsx')

upload_doc = st.file_uploader("Upload your approved Approval of Requirement (AOR) here. This document should include information on your requirement", type="docx",key="file_uploader_1")
upload_pdf= st.file_uploader("Upload your GeBIZ Summary of Offer here. This document should indicate the summary of prices offered by tenderers.", type="pdf",key="file_uploader_2")
upload_excel = st.file_uploader("Uploade your Detailed Evaluation Report here. This document should contain the evaluation results for each offer, including PQ score and brief assessment., including scores and comments.", type="xlsx", key="file_uploader_3")

if st.button("Process Now!") and upload_excel is not None and upload_doc is not None and upload_pdf is not None:
    
    with st.spinner('Loading...'):
        docx_bytes = upload_doc.read()
        docxfile = io.BytesIO(docx_bytes)
        doc = Document(docxfile)

        pdf_bytes = upload_pdf.read()
        pdffile = io.BytesIO(pdf_bytes)
        pdf_text= read_pdf_file(pdffile)
        data = read_excel_data(upload_excel)
        doctext = ""
        for paragraph in doc.paragraphs:
            doctext += paragraph.text + "\n"

        table_string =""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_string += cell.text + "\t"  # Add tab separator between cells
                table_string += "\n"  # Add newline after each row

        #create new document
        document = Document()

        st.write("Give us a moment while we process your data.")
        set_document_font(document, "Arial", 12)

        #retrieve the Provision of xxx information
        prompt = f"""
            
            Search for the description and retrieve only the sentence for me in string (For example, Provision of Comprehensive XXX for Infrastructure System at <Location>).
            
            Please ensure you only take records found in the document.
            `{table_string}`
            """
        response = get_completion(prompt)
        subject = response.strip('"')

        prompt = f"""
            
            I only need the following paragraph for my prompt response and this should include the following:

            Award Recommendation for the <<Subject>> (USER TO INSERT AGY2024TDR002)
            * Search for the Subject and replace the <<Subject>> in the above paragraph(For example, Provision of XXX Services to Information Technology (IT) Infrastructure System at Government Office)

            Please ensure you only take records found in the document.
            `{doctext}`
            """
        response = get_completion(prompt)
        table = document.add_table(rows=3, cols=3)

        widths = (Cm(5),Cm(1),Cm(10))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        table.cell(0, 0).text = "Subject"
        table.cell(0, 1).text = ":"
        table.cell(0, 2).text = response

        table.cell(1, 0).text = "For"
        table.cell(1, 1).text = ":"
        table.cell(1, 2).text = "Approval"

        table.cell(2, 0).text = "No. of Pages"
        table.cell(2, 1).text = ":"
        table.cell(2, 2).text = "<<PLEASE INSERT NO. OF PAGES"

        document.add_paragraph(" ")

        prompt = f"""

                From the table under the offer summary, extract the following information for each record. 
                
                Return only Supplier name with the lowest Total Responded Amount in string, capital letter for start of each word.
                **Avoid including any additional text or headers.**

                `{pdf_text}`
                """
        response = get_completion2(prompt)
        suppliername = response.strip()
        #check for provisional sum
        prompt = f"""
            Search for the percentage amount that is within Provisional Sum or Contingency Sum. For example, Provisional Sum (~10% of the awarded). return only the number if you can find contingency or provisional sum.
            `{doctext}`
            """
        response = get_completion(prompt)
        storeprovisionalamount = int(response) + 0

        if storeprovisionalamount>0:

            prompt = f"""

                From the table under the offer summary, extract the following information for each record. 
                
                Return only lowest Total Responded Amount in integer format.
                **Avoid including any additional text or headers.**

                `{pdf_text}`
                """
            response = get_completion2(prompt)
            lowestapvtotal = int(response.strip())

        lowestapvtotalalone = lowestapvtotal
        lowestapvtotal = ((storeprovisionalamount/100) * lowestapvtotal) + lowestapvtotal
        lowestapvtotalint = int(lowestapvtotal)

        lowestapvtotal = round_to_nearest_thousand(lowestapvtotal)

        if lowestapvtotalint < 1000000:
            selected_option = "A"
        elif lowestapvtotalint >100000 and lowestapvtotalint <10000000:
            selected_option = "B"
        else:
            selected_option = "C"

        #aim
        prompt = f"""
            
            I only need the following paragraph for my prompt response and this should include the following:

            Tender Board {selected_option} 's Approval is sought to award the open/closed tender to {suppliername}  for the <<Subject>> at an Approved Procurement Value (APV) of $`{lowestapvtotal}` 
            
            Search for the Subject and replace the <<Subject>> in the above paragraph(For example, Provision of Comprehensive Maintenance Services to Information Technology (IT) Infrastructure System at Government Office)
            * All amounts in integer format (e.g. 100,000), no decimal points. 

            Please ensure you only take records found in the document.
            `{doctext}`
            """
        response = get_completion(prompt)
        #heading_1= document.add_heading("A. Aim", level=1)
        force_font("A.  Aim") #Forces the font to be bold and arial
        #st.write(response)

        document.add_paragraph(response)

        #Background
        prompt = f"""
            I only need the following paragraph for my prompt response and this should include the following:

            The requirement for <<Subject>> for a <<period>> has been approved by <<Approver>> on <<USER TO INSERT DATE>> at an Estimated Procurement Value(EPV) of 
            <<EPV>> (excluding GST)(see AOR Paper<<USER TO INSERT AOR PAPER FROM MPS>>)

            * Subject is Search for the Subject and replace the <<Subject>> in the above paragraph(For example, Provision of Comprehensive Maintenance Services to Information Technology (IT) Infrastructure System at Government Office)
            * Search for Period and replace the <<period>> in the paragraph above. Period will be found under the Aim section which usually states period of X(a number) of Months(e.g. firm period of 24 months)
            * Search for the Estimated Procurement Value (EPV) and replace the <<EPV>> in the paragraph above with the amount.
            * Search for the EPV in the document and determine the approve from the list below. Approvers limits will have two levels, one is Project AOR and the other is Non-Project AOR, the approver limits include:
                Select [Project AOR Approver] based on the following information:
                - Project cost from $6000 to 25,000: Director (D)
                - Project cost from $25001 to $90,000: Group Director(GD)
                - Project cost from $90,001 to $50000: GD co-opted into EXCO
                - Project cost from $500,001 to $1,000,000: Deputy Chief Executive Officer(DCE) and any one Group Director (GD)
                - Project cost from more than $10 million to $20 milllion: Chairman and Chief Executive Officer(CEO)
                - Project cost above $20 million: CPF Board
                Select[non-project AOR Approver] based on the following information:
                - Estimated Procurement Value or EPV from $6000 to $25000 is Director (D)
                - Estimated Procurement Value or EPV from $6000 to 25,000: Director (D)
                - Estimated Procurement Value or EPV from $25001 to $90,000: Group Director(GD)
                - Estimated Procurement Value or EPV from $90,001 to $50000: GD co-opted into EXCO
                - Estimated Procurement Value or EPV from $500,001 to $1,000,000: Deputy Chief Executive Officer(DCE) and any one Group Director (GD)
                - Estimated Procurement Value or EPV from more than $10 million: Chairman and Chief Executive Officer(CEO)

            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
            `{doctext}`
            """
        response = get_completion(prompt)
        #heading_1= document.add_heading("B. Background")
        force_font("B.  Background") #Forces the font to be bold and arial
        document.add_paragraph(response)

        #Procurement Approach
        prompt = f"""
            I have a PDF document containing an offer summary. Please extract the following information into this paragraph.
            
            I only need the following paragraph for my prompt response and this should include the following:

            An open tender was published in GeBIZ on <<Publication Date>> and closed on <<Closing Date>>
            * Publication date will be placed into <<Publication Date>>
            * Closing date and time will be placed into <<Closing Date>>. Do not need to include time.

            `{pdf_text}`
            """
        response = get_completion2(prompt)
        #st.write(response)
        #heading_1= document.add_heading("C. Procurement Approach")
        force_font("C.  Procurement Approach") #Forces the font to be bold and arial
        document.add_paragraph(response)

        #Tender Evaluation Committee

        prompt = f"""
            I only need the following paragraph for my prompt response and this should include the following:

            * Subject is Search for the Subject and replace the <<Subject>> in the above paragraph(For example, Provision of Comprehensive Maintenance Services to Information Technology (IT) Infrastructure System at Government Office)
            * Search for Period and replace the <<period>> in the paragraph above. Period will be found under the Aim section which usually states period of X(a number) of Months(e.g. firm period of 24 months)
            * Search for the Estimated Procurement Value (EPV) and replace the <<EPV>> in the paragraph above with the amount.
            * Search for the EPV in the document and determine the approve from the list below. Approvers limits will have two levels, one is Project AOR and the other is Non-Project AOR, the approver limits include:
                Select [Project AOR Approver] based on the following information:
                - Project cost from $6000 to 25,000: Director (D)
                - Project cost from $25001 to $90,000: Group Director(GD)
                - Project cost from $90,001 to $50000: GD co-opted into EXCO
                - Project cost from $500,001 to $1,000,000: Deputy Chief Executive Officer(DCE) and any one Group Director (GD)
                - Project cost from more than $10 million to $20 milllion: Chairman and Chief Executive Officer(CEO)
                - Project cost above $20 million: CPF Board
                Select[non-project AOR Approver] based on the following information:
                - Estimated Procurement Value or EPV from $6000 to $25000 is Director (D)
                - Estimated Procurement Value or EPV from $6000 to 25,000: Director (D)
                - Estimated Procurement Value or EPV from $25001 to $90,000: Group Director(GD)
                - Estimated Procurement Value or EPV from $90,001 to $50000: GD co-opted into EXCO
                - Estimated Procurement Value or EPV from $500,001 to $1,000,000: Deputy Chief Executive Officer(DCE) and any one Group Director (GD)
                - Estimated Procurement Value or EPV from more than $10 million: Chairman and Chief Executive Officer(CEO)

            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
            `{doctext}`
            """
        response = get_completion(prompt)
        #st.write(response)
        heading_1= document.add_heading("D. Tender Evaluation Committee", level=1)
        document.add_paragraph("A Tender Evaluation Committee (TEC), comprising the following officers, evaluated the proposal")
        document.add_paragraph("Table 1: Composition of Tender Evaluation Committee")
        table = document.add_table(rows=4, cols=4)
        table.style = 'Table Grid'


        widths = (Cm(1),Cm(6),Cm(6),Cm(3.7))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        table.cell(0, 0).text = "S/N"
        table.cell(0, 1).text = "Name"
        table.cell(0, 2).text = "Designation"
        table.cell(0, 3).text = "Role"

        set_cell_colour(table) # Set the header row background to grey

        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "<<INSERT NAME>>"
        table.cell(1, 2).text = "<<INSERT DESIGNATION>>"
        table.cell(1, 3).text = "Chairperson"

        table.cell(2, 0).text = "2"
        table.cell(2, 1).text = "<<INSERT NAME>>"
        table.cell(2, 2).text = "<<INSERT DESIGNATION>>"
        table.cell(2, 3).text = "Member"

        table.cell(3, 0).text = "3"
        table.cell(3, 1).text = "<<INSERT NAME>>"
        table.cell(3, 2).text = "<INSERT DESIGNATION>>"
        table.cell(3, 3).text = "Member"

        document.add_paragraph("The TEC as listed above declared that they have no vested interest in the procurement and no conflict of interest with any of the suppliers participating in this procurement (see PROC/COI/20XX/XXXXX for the declaration for conflict of interest by the TEC.)")

        #Summary of Offers Section
        force_font("E.  Offers Received") #Forces the font to be bold and arial
        prompt = f"""
            Form the sentence below:
            A total of <Offer Numbers> were received as summarised below (See Annex X for Details)
            From the table under the offer summary, Can you only return me the number of rows contained in the records and fill in the numbers in <Offer Numbers>

            `{pdf_text}`
            """
        response = get_completion2(prompt)
        document.add_paragraph(response)

        response = get_completion2(prompt)

        document.add_paragraph("Table 2: Summary of Offers")

        prompt = f"""

            From the table under the offer summary, extract the following information for each record. 
            
            Return only in string format the Supplier name and Total Responded Amount with a delimiter of ":" followed by a ; to denote next company name. For example (Company A: $100,000, Comapny B: $120000)

            * Supplier name
            * Total responded amount in integer format (e.g. 100000), no decimal points. 

            Please sort the data by ascending order by on Total Responded amount. 
            Capitalise each word for supplier name.
            **Avoid including any additional text or headers.**

            `{pdf_text}`
            """
        response = get_completion2(prompt)
        storesupplieramount = response
        # Clean and split the text by rows
        cleaned_text = response.strip()
        rows = cleaned_text.split(";")

        #count no. of rows in records.
        prompt = f"""
            From the table under the offer summary, return only number of rows in integer format. I am specifically asking for only a number digit.
            `{pdf_text}`
            """
        rownum = get_completion2(prompt)
        tablenorow = int(rownum) + 1
        table = document.add_table(rows=int(tablenorow), cols=3)
        table.style = 'Table Grid'
        widths = (Cm(1),Cm(8),Cm(5))

        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # Extract company name and amount from each entry
        for row_index, row in enumerate(rows, start=1):
            columns = row.split(":")


            # Insert serial number in the first column
            cell = table.cell(row_index, 0)
            cell.text = str(row_index)  # Insert the serial number
            
            # Insert data into the remaining columns
            for col_index, column in enumerate(columns, start=1):  # Data starts from col 1
                if col_index == 1:
                    column = column.title()
                elif col_index == 2:
                    column = "${:,}".format(int(column))
                cell = table.cell(row_index, col_index)
                cell.text = column.strip()

        table.cell(0, 0).text = "S/N"
        table.cell(0, 1).text = "Description"
        table.cell(0, 2).text = "Total Price(S$)"
        set_cell_colour(table)

        #Evaluation Section
        #heading_1= document.add_heading("F. Evaluation", level=1)
        force_font("F.  Evaluation")#Forces the font to be bold and arial

        #Produce Table 3: Summary of Scores

        # Analyze the data
        prompt = f"""
            You are looking through an excel data sheet that comprises of S/N, Evaluation Criteria and Component/Score(e.g. Compliant). This are followed by all the columns that contain the company names and their respective scoring and tender amount.
        
            Drill down to the Price Score Section and help to present the dataset so i can use to parse in a word document table later.
            Each start of the row, delimiter with a ":" will contain company name, method statement points, quality of proposed manpower, track records, performance evaluation, quality score, price score, total score
            Each row of data will have a delimiter of ";" 
            
            locate the following:
            *Locate the company name
            *locate Method Statement Points
            *Locate QUality of Proposed Manpower
            *Locate Track Records
            *Locate Performance Evaluation
            *Quality score is sum of computation of Method Statement Points, Quality of Proposed Manpower, track records and performance evaluation.
            *Locate Price Score
            *Total Score is the Sum of Price Score and Quality Score. 

            **Ignore the first two rows**
            **The 2nd row is the dataframe headers**
            **Avoid including any additional text or headers.**
            **All numbers with decimal places will be rounded to two decimal places.**
            **Remove " in the response**
            Please ensure you only take records found in the document.

            `{data}`
            """

        response = get_completion2(prompt)
        cleaned_dataset = response.strip("`")

        rows = cleaned_dataset.split(";")
        cleaned_data = [entry for entry in rows if entry.strip()]
        #st.write(cleaned_data)
        prompt = f"""
            Summarise the data in terms of Quality for the top 3 companies only, mentioned who is ranked 1st. 
            Each row according to the delimiter of ":", are providing the following information:company name, method statement points, quality of proposed manpower, track records, performance evaluation, quality score, price score, total score

            Could you breakdown according to quality of proposed manpower, track records, method statement points for each company and provide me 3 paragraphs. Be as concise and formal in the writing.
            
            *Capitalise fhe first letter for each word for company name.
            Please ensure you only take records found in the document.
            
            `{cleaned_data}`
            """
        response = get_completion2(prompt)
        cleaned_text = response.strip('"')
        #st.write(cleaned_text)
        document.add_paragraph(cleaned_text)

        #count no. of rows in records.
        prompt = f"""
            From the table under the offer summary, return only number of rows in integer format. I am specifically asking for only a number digit.
            `{cleaned_data}`
            """
        rownum = get_completion2(prompt)
        #st.write(rownum)
        document.add_paragraph("Table 3: Summary of Scores")

        tablenorow = int(rownum) + 1
        table = document.add_table(rows=int(tablenorow), cols=9)
        table.style = 'Table Grid'
        widths = (Cm(2),Cm(5),Cm(5),Cm(4),Cm(4),Cm(3),Cm(4),Cm(4),Cm(5))

        for row_index, row in enumerate(cleaned_data, start=1):  # Start at 1 for serial numbers
            columns = row.split(":")
            
            # Insert serial number in the first column
            cell = table.cell(row_index, 0)
            cell.text = str(row_index)  # Insert the serial number
            
            # Insert data into the remaining columns
            for col_index, column in enumerate(columns, start=1):  # Data starts from col 1
                cell = table.cell(row_index, col_index)

                # Capitalize first letter of each word in the "Tenderer" cell
                if col_index == 1:  # Check if it's the "Tenderer" column
                    cell.text = column.strip().title()
                else:
                    cell.text = column.strip()

        table.cell(0, 0).text = "S/N"
        table.cell(0, 1).text = "Tenderer"
        table.cell(0, 2).text = "Method Statement (5pts)"
        table.cell(0, 3).text = "Quality of Proposed Manpower (5pts)"
        table.cell(0, 4).text = "Track records (10pts)" 
        table.cell(0, 5).text = "Performance Evaluation (10pts)"
        table.cell(0, 6).text = "[A] Total Quality (30pts)"
        table.cell(0, 7).text = "[B] Price (70pts)"	
        table.cell(0, 8).text = "[A] + [B] Total PQ Score"
        set_cell_colour(table)

        force_font("G. Price Assessment") #Forces the font to be bold and arial
        #first paragraph

        string_with_rows=""
        for row in rows:
            string_with_rows += row + "\n" 
            combined_text = doctext + " " + storesupplieramount

        #calculate APV
        prompt = f"""
            Please anaylse the data found in the Annexes. Please return only the lowest total responded amount as an integer value.
            **Avoid including any additional text or headers.**
            
            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
            `{combined_text}`
            """
            
        response = get_completion(prompt)
        cleaned_text = response.strip()
        apv_sum = int(cleaned_text) 
        #elaborate on Price Assessment
        #st.write(data)

        #prompt = f"""
        #    You are looking through an excel data sheet that comprises of S/N, Evaluation Criteria and Component/Score(e.g. Compliant). This are followed by all the columns that contain the company names and their respective scoring and tender amount.
        
        #    Drill down to the Price Score Section and locate the following:
        #    Explain to me how many companies can you see from the dataframe.
        #    *Locate and Include the Price($) under each company column. 
        #    *Locate and include Total score.
        #    **Ignore the first two rows**
        #    **The 2nd row is the dataframe headers**
        #    Please ensure you only take records found in the document.
        #    `{data}`
        #    """

        #cleaned_dataset = get_completion2(prompt)

        #st.write(cleaned_dataset)
        prompt = f"""
            Imagine you are a Tender Evaluation Committee(TEC) Member. Mention you did due diligence. Write the given pointers in 3 paragraphs. Be as formal and concise in 3rd person writing.
            in this dataset, each row according to the delimiter of ":", are providing the following information:company name, method statement points, quality of proposed manpower, track records, performance evaluation, quality score, price score, total score

            *Under the Price Competitiveness, highlight how many proposals did Tender Evaluation Committee(TEC) evaluate in detail and why?
            *Give a summary of overall scores presented in the Total Score.
            *Breakdown which company submitted the lowest-priced tender, and what was their Price Quality(PQ) Score. PQ Score is the Total Score at the last row, evaluate who has the highest total score.
            
            **Avoid including any additional text or headers**
            **Avoid all Markdown format**
            *Capitalise the first letter of each word of the company name.

            Please ensure you only take records found in the document.
            `{cleaned_data}`
            """
        response = get_completion2(prompt)
        cleaned_text = response.strip('"')

        #st.write(cleaned_text)
        document.add_paragraph(cleaned_text)
        #document.add_paragraph(response)

        #calculate EPV no provisional sum
        prompt = f"""
            return only the value next to the description that states Provision of XXX in integer format.
            **Avoid including any additional text or headers.**

            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
            `{table_string}`
            """
        response = get_completion(prompt)
        cleaned_text = response.strip()
        epvnops_sum = int(cleaned_text) 

        #calculate epv_difference
        epv_difference = epvnops_sum- apv_sum

        prompt = f"""
            I only need the following paragraph for my prompt response, you do not need to tell me Proposed Paragraph, and this should include the following:
            
            1. <Lowest Quoted Supplier Price Name> is assessed to be reasonable because their quotes is the lowest among the tenderers and it is also lower than the EPV of $`{epvnops_sum}` by $`{epv_difference}`
                *Derive the <Lowest Quoted Supplier Price Name> based on the total responded amount. Capitalise each word.

            **Avoid adding Angular Brackets(e.g. <>) in the Lowest Quoted Supplier Price Name.
            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
            `{combined_text}`
            """
        response = get_completion(prompt)
        document.add_paragraph(response)

        #Recommendation Section
        #heading_1= document.add_heading("H. Recommendation", level=1)
        force_font("H.  Recommendation") #Forces the font to be bold and arial

        prompt = f"""
            I only need the following paragraph for my prompt response, you do not need to tell me Proposed Paragraph, and this should include the following:
                
            1.  In view of the above, TEC recommends to award the open tender to {suppliername} at a total APV of $`{lowestapvtotal}` with breakdown as shown in table 2 below.
                    Capitalise the first letter of each word for supplier name.

            **derive the amounts such as $ with a comma in thousands (e.g. 100,000), no decimal points needed**
            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
            `{combined_text}`
            """

        response = get_completion(prompt)
        document.add_paragraph(response)
        document.add_paragraph("Table 2: Breakdown of APV")

        prompt = f"""
            Provide Only me a tuple list with delimiter ";" <Subject>, <EPV>, <Difference>. Follow strictly to the order. 
            *Only information provided in <>. 
            *Remove backticks(e.g. '''') and square brackets (e.g. []) and brackets (e.g. () )
                
            *derive it subject and store it as <Subject>, phrase it as Provision of (services) at (location)
            *Search for the Estimated Procurement Value (EPV) and replace the <EPV> in the paragraph above with the amount.
            *Derive a <Difference> value by subtracting the EPV value by the lowest supplier price. provide the value in dollar format.

            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.

            `{combined_text}`
            """

        response = get_completion(prompt)
        cleaned_text = response.strip()
        parsed_response = response.split(";")
        tuple_list = []
        for item in parsed_response:
            tuple_list.append(tuple(item.split(",")))

        #check for provisional sum
        
        prompt = f"""
            Search for the information if it contains Provisional Sum or Contingency Sum. For example, Provisional Sum (~10% of the awarded).
            if you can find contingency or provisional sum, return one value true.
            if not, return false.
            `{combined_text}`
            """

        #sett table size row to 5 if there is provisional sum or contingency sum
        provisionaltablesize = 4
        response = get_completion(prompt)
        #response = get_completion(prompt)
        #st.write(response)
        if response == "True":
            provisionaltablesize = 5

        #st.write(provisionaltablesize)
        table = document.add_table(rows=provisionaltablesize, cols=3)
        table.style = 'Table Grid'
        set_cell_colour(table)
        widths = (Cm(1),Cm(11),Cm(4))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        table.cell(0, 0).text = "S/N"
        table.cell(0, 1).text = "Description"
        table.cell(0, 2).text = "Total Price ($)"
        set_cell_colour(table) # Set the header row background to grey

        #awarded row
        table.cell(1, 0).text = "Awarded"
        row_cells = table.rows[1].cells
        row_cells[0].merge(row_cells[2])

        #1st row showing the 1st item
        table.cell(2, 0).text = "1"
        table.cell(2, 1).text = subject
        #table.cell(2, 2).text = "${:,}".format(apv_sum)
        table.cell(2, 2).text = "$"+ round_to_nearest_thousand(lowestapvtotalalone)
        #st.write(response)

        if response == "True":

            #calculate provisional/contingency sum percentage
            prompt = f"""
                return only the percentage in the provisional or contingency sum information in integer, no % needed.
                Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.
                `{table_string}`
            """
            response = get_completion(prompt)
            cleaned_text = response.strip()
            provisional_sum_percent = int(cleaned_text) /100

            #response = get_completion(prompt)
            #cleaned_text = response.strip()
            provisional_sum = round(int(apv_sum) * provisional_sum_percent)    
            
            #check for provisional sum
            prompt = f"""
                Search for the information if it contains Provisional Sum or Contingency Sum. For example, Provisional Sum (~10% of the awarded). return the phrase provisional sum or contingency sum if you can find contingency or provisional sum.
                If you cannot find, just reply none
            `{table_string}`
            """
            provisionalsumtext = get_completion(prompt)
            cleaned_text = provisionalsumtext.strip()
            #st.write(provisional_sum_percent)

            if provisionalsumtext =="provisional sum":
                percentsum ="Provisional Sum ("+ str(provisional_sum_percent* 100) + "% of the Awarded)"
                table.cell(3, 0).merge(table.cell(3, 1))  # Merge cells in the third row
                table.cell(3, 1).text = percentsum
            elif provisionalsumtext == "contingency sum":
                percentsum ="Contingency Sum ("+ str(provisional_sum_percent * 100) + "% of the Awarded)"
                table.cell(3, 0).merge(table.cell(3, 1))  # Merge cells in the third row
                table.cell(3, 1).text = percentsum
            else:
                percentsum ="none"
            
            table.cell(3, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(3, 1).paragraphs[0].runs[0].bold = True

            table.cell(3, 2).text = "$" +round_to_nearest_thousand(provisional_sum)

            table.cell(4, 0).text = "Approved Procurement Value (APV)"

            table.cell(4, 0).merge(table.cell(4, 1))  # Merge cells in the third row
            table.cell(4, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(4, 0).paragraphs[0].runs[0].bold = True

            #table.cell(4, 2).text = "${:,}".format(round(provisional_sum +apv_sum))
            table.cell(4, 2).text = "$" + lowestapvtotal
            prompt = f"""
            I only need the following paragraph. 
            The APV is within the approved EPV of <EPV value>. The provisional sum is set aside for the additional goods or services related to the scope of the Contract.
            * Obtain and replace the <EPV value> into the above paragraph.
            Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.

            `{doctext}`
            """
        else:
            table.cell(3, 0).text = "Approved Procurement Value (APV)"
            table.cell(3, 0).merge(table.cell(3, 1))  # Merge cells in the third row
            table.cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.cell(3, 0).paragraphs[0].runs[0].bold = True
            table.cell(3, 2).text = "$" + lowestapvtotal

            prompt = f"""
                I only need the following paragraph. 
                The APV is within the approved EPV of <EPV value>.
                * Obtain and replace the <EPV value> into the above paragraph.
                Please ensure that all records are extracted, even if they are in a non-standard format. If you encounter any difficulties, please provide specific details about the issues you're facing.

                `{doctext}`
            """

        response = get_completion(prompt)
        document.add_paragraph(response)

        #Recommendation Section
        #heading_1= document.add_heading("I. Approval",level=1)
        force_font("I.  Approval") #Forces the font to be bold and arial
        prompt = f"""
            Provide me only the following paragraph without telling me my prompt response and any other responses
            In accordance with CPFB's Financial Regulations Section <Section>, Tender Board {selected_option}'s Approval is sought to award the open tender to {suppliername} for the {subject} at an Approved Procurement (APV) of $`{lowestapvtotal}`(excluding GST).
            * Find the project type under subject and fill up <Section>.
            - If Project AOR, refer to 4.2.3
            - if non-project AOR, refer to 4.2.2. 

            `{doctext}`
            """

        response = get_completion(prompt)
        response = response.strip()
        document.add_paragraph(response)
        #st.write("Approval segment")
        #st.write(response)

        prompt = f"""
            Locate and provide only the name, designation and department who prepared the document.
            his information can be found under <Prepared By> at the segment of the document. Provide a line break after each segment. 
            If there are more than one name, please break the name into the next line, followed by department, and the next record in the name and department.
            They are not the DCE or the D of the department.

            `{table_string}`
            """
        response = get_completion(prompt)

        table = document.add_table(rows=3, cols=3)
        widths = (Cm(4),Cm(0.8),Cm(10))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        table.cell(0, 0).text = "Prepared By"
        table.cell(0, 1).text = ":"
        table.cell(0, 2).text = response

        prompt = f"""
            I only need the following information on the name of the person, followed by a line break and provide its designation and department who provided inputs to the document. 
            his or her name can be found under "With Inputs From". The person is also usually from Procurement department.
            If there are more than one name, please break the name into the next line, followed by department, and the next record in the name and department.
            Example: Tan Chin Soo
                    Assistant Manager, Procurement Management

                    John Henry
                    Assistant Director, Procurement Management
            `{table_string}`
            """
        response = get_completion(prompt)

        table.cell(1, 0).text = "With Inputs From"
        table.cell(1, 1).text = ":"
        table.cell(1, 2).text = response

        #formatting of document
        set_all_headings_bold(document)
        for run in heading_1.runs:
            run.font.name = 'Arial'  # Change the font to Arial
            run.font.size = Pt(12)   # Optional: Change font size

        #finally save the document
        output_file = "ouputfile.docx"
        document.save(output_file)
        st.success("Documents merged and saved successfully!Click the Button Below to download it.")

        # Download the saved document
        with open(output_file, "rb") as f:
            st.download_button(
                label="Download Processed Document",
                data=f.read(),
                file_name="processed_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.warning("Please upload all documents.")
